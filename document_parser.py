import os
from pathlib import Path
import PyPDF2
import docx
from typing import List, Dict
import streamlit as st
import tabula
import pandas as pd


class DocumentParser:
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt']
    
    def parse_pdf_with_tables(self, file_path: str) -> Dict:
        """Extract text and tables from PDF files using tabula and PyPDF2"""
        text_content = ""
        table_content = ""
        
        # Extract text using PyPDF2
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
        except Exception as e:
            print(f"Error extracting text: {e}")
        
        # Extract tables using tabula
        try:
            tables = tabula.read_pdf(file_path, pages='all', multiple_tables=True)
            for i, table in enumerate(tables):
                table_content += f"\n--- Table {i+1} ---\n"
                table_content += table.to_string() + "\n"
        except Exception as e:
            print(f"Error extracting tables: {e}")
        
        return {
            'text_content': text_content,
            'table_content': table_content,
            'combined_content': text_content + table_content
        }
    
    def parse_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        result = self.parse_pdf_with_tables(file_path)
        return result['combined_content']
    
    def parse_docx(self, file_path: str) -> str:
        """Extract text from Word documents"""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract tables from docx
        for table in doc.tables:
            text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                text += " | ".join(row_text) + "\n"
        
        return text
    
    def parse_txt(self, file_path: str) -> str:
        """Extract text from TXT files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def parse_document(self, file_path: str) -> Dict:
        """Parse any supported document format"""
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == '.pdf':
            content = self.parse_pdf(file_path)
        elif file_ext == '.docx':
            content = self.parse_docx(file_path)
        elif file_ext == '.txt':
            content = self.parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        return {
            'filename': Path(file_path).name,
            'content': content,
            'file_type': file_ext
        }
